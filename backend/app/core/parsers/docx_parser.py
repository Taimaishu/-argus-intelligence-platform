"""DOCX document parser using python-docx."""

from docx import Document
from .base import BaseParser, ParsedDocument


class DOCXParser(BaseParser):
    """Parser for DOCX documents."""

    def supports_file_type(self, file_extension: str) -> bool:
        """Check if file extension is .docx"""
        return file_extension.lower() == ".docx"

    def parse(self, file_path: str) -> ParsedDocument:
        """
        Parse DOCX document and extract text and metadata.

        Args:
            file_path: Path to DOCX file

        Returns:
            ParsedDocument with extracted content
        """
        try:
            doc = Document(file_path)

            # Extract metadata from core properties
            core_props = doc.core_properties
            title = core_props.title
            author = core_props.author

            # Extract text from paragraphs
            full_text = ""
            sections = []
            current_section = {"heading": "Introduction", "content": ""}

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()

                if not text:
                    continue

                # Check if paragraph is a heading (style-based)
                if paragraph.style.name.startswith("Heading"):
                    # Save previous section if it has content
                    if current_section["content"]:
                        sections.append(current_section.copy())

                    # Start new section
                    current_section = {"heading": text, "content": ""}
                else:
                    current_section["content"] += text + "\n"

                full_text += text + "\n"

            # Add last section
            if current_section["content"]:
                sections.append(current_section)

            # Extract text from tables
            for table in doc.tables:
                table_text = "\n"
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    table_text += row_text + "\n"
                full_text += table_text + "\n"

            return ParsedDocument(
                text=self.clean_text(full_text),
                title=title,
                author=author,
                metadata={
                    "subject": core_props.subject,
                    "created": str(core_props.created) if core_props.created else None,
                    "modified": str(core_props.modified) if core_props.modified else None,
                    "keywords": core_props.keywords,
                },
                sections=sections
            )

        except Exception as e:
            raise Exception(f"Failed to parse DOCX: {str(e)}")
